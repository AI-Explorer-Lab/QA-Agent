from __future__ import annotations

import argparse
import re
import sys
from pathlib import Path
from typing import Dict, Iterator, List

from docx import Document

PROJECT_ROOT = Path(__file__).resolve().parents[1]
if str(PROJECT_ROOT) not in sys.path:
    sys.path.insert(0, str(PROJECT_ROOT))

from chunking_service.structured_chunking import wrap_table_markdown  # noqa: E402


def _default_docx_path() -> Path:
    test_docs_dir = PROJECT_ROOT / "docs" / "test_docs"
    candidates = sorted(test_docs_dir.glob("MinerU_docx_*.docx"))
    if candidates:
        return candidates[-1]
    return test_docs_dir / "sample.docx"


DEFAULT_DOCX_PATH = _default_docx_path()
DEFAULT_OUTPUT_PATH = PROJECT_ROOT / "test" / "docx_body_ptbl_processed.txt"


def _normalize_text(text: str) -> str:
    return re.sub(r"\s+", " ", text or "").strip()


def iterate_processed_ptbl(doc) -> Iterator[Dict]:
    """
    Keep traversal logic consistent with project chunking code:
    loop over doc.element.body and match paragraph/table by _element identity.
    Output only processed paragraph/table payloads.
    """
    table_index = 0

    for idx, element in enumerate(doc.element.body):
        if element.tag.endswith("p"):
            for para in doc.paragraphs:
                if para._element == element:
                    text = _normalize_text(para.text)
                    if text:
                        yield {
                            "index": idx,
                            "kind": "p",
                            "content": text,
                        }
                    break
            continue

        if element.tag.endswith("tbl"):
            for table in doc.tables:
                if table._element == element:
                    table_index += 1
                    rows: List[List[str]] = []
                    for row in table.rows:
                        cells = [_normalize_text(cell.text) for cell in row.cells]
                        if any(cells):
                            rows.append(cells)

                    if rows:
                        yield {
                            "index": idx,
                            "kind": "table",
                            "table_id": f"table_{table_index}",
                            "content": wrap_table_markdown(rows),
                        }
                    break


def dump_processed_ptbl(docx_path: Path, output_path: Path) -> int:
    if not docx_path.exists():
        raise FileNotFoundError(f"docx not found: {docx_path}")

    doc = Document(str(docx_path))
    records = list(iterate_processed_ptbl(doc))
    output_path.parent.mkdir(parents=True, exist_ok=True)

    blocks: List[str] = []
    for seq, item in enumerate(records, start=1):
        if item["kind"] == "p":
            block = f"[{seq}] p\n{item['content']}"
        else:
            block = f"[{seq}] table ({item['table_id']})\n{item['content']}"
        blocks.append(block)

    output_path.write_text("\n\n".join(blocks), encoding="utf-8")
    return len(records)


def build_arg_parser() -> argparse.ArgumentParser:
    parser = argparse.ArgumentParser(
        description="Dump processed p/tbl blocks in body order (one block per paragraph/table, blank-line separated)."
    )
    parser.add_argument("--docx", type=str, default=str(DEFAULT_DOCX_PATH))
    parser.add_argument("--output", type=str, default=str(DEFAULT_OUTPUT_PATH))
    return parser


if __name__ == "__main__":
    args = build_arg_parser().parse_args()
    docx_path = Path(args.docx)
    output_path = Path(args.output)

    count = dump_processed_ptbl(docx_path=docx_path, output_path=output_path)
    print(f"processed_blocks={count}")
    print(f"output={output_path}")
